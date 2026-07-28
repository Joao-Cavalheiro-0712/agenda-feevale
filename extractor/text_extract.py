"""Extrai texto bruto de PDF, Word (.docx/.doc), Excel e CSV/TXT."""
from __future__ import annotations

import csv
import io
import os
import re
import shutil
import subprocess
import tempfile


def extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return _from_pdf(data)
    if ext == ".docx":
        return _from_docx(data)
    if ext == ".doc":
        return _from_doc(data)
    if ext in (".xlsx",):
        return _from_xlsx(data)
    if ext in (".csv",):
        return _from_csv(data)
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError(
        f"Formato não suportado: {ext or '(sem extensão)'}. "
        "Use PDF, DOC, DOCX, XLSX, CSV ou TXT."
    )


def _from_pdf(data: bytes) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            parts.append(txt)
            # Tabelas costumam guardar cronogramas; achatamos em linhas.
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [c for c in row if c]
                    if cells:
                        parts.append(" | ".join(str(c) for c in cells))
    return "\n".join(parts)


def _from_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_doc(data: bytes) -> str:
    """Lê o formato binário antigo do Word (.doc)."""
    # Alguns arquivos .doc na verdade são .docx (zip começa com 'PK').
    if data[:2] == b"PK":
        return _from_docx(data)

    # 1) antiword (ferramenta dedicada; instalada via Nixpacks no Railway).
    text = _antiword(data)
    if text and text.strip():
        return text

    # 2) fallback em Python puro via olefile (melhor esforço).
    text = _doc_fallback(data)
    if text and text.strip():
        return text

    raise ValueError(
        "Não consegui ler este .doc. Reabra no Word e salve como .docx "
        "(ou exporte em PDF) e envie novamente."
    )


def _antiword(data: bytes) -> str:
    if not shutil.which("antiword"):
        return ""
    tmp = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(data)
            tmp = f.name
        out = subprocess.run(
            ["antiword", "-m", "UTF-8.txt", "-w", "0", tmp],
            capture_output=True,
            timeout=60,
        )
        if out.returncode == 0:
            return out.stdout.decode("utf-8", errors="ignore")
    except (subprocess.SubprocessError, OSError) as exc:
        print(f"[text_extract] antiword falhou: {exc}")
    finally:
        if tmp and os.path.exists(tmp):
            os.unlink(tmp)
    return ""


def _doc_fallback(data: bytes) -> str:
    """Extrai texto legível do stream WordDocument sem dependência externa."""
    try:
        import olefile
    except ImportError:
        return ""
    if not olefile.isOleFile(io.BytesIO(data)):
        return ""
    try:
        with olefile.OleFileIO(io.BytesIO(data)) as ole:
            if not ole.exists("WordDocument"):
                return ""
            raw = ole.openstream("WordDocument").read()
    except Exception:  # noqa: BLE001
        return ""

    # O texto costuma estar em cp1252/utf-16 misturado com bytes de controle.
    txt = raw.decode("cp1252", errors="ignore")
    # Mantém apenas trechos imprimíveis (letras, números, pontuação comum).
    txt = re.sub(r"[^\x09\x0a\x0d\x20-\x7e -ɏ]", " ", txt)
    txt = re.sub(r"[ \t]{2,}", " ", txt)
    # Descarta linhas curtas/ruído, preserva as que parecem conteúdo real.
    linhas = [
        ln.strip()
        for ln in re.split(r"[\r\n]+", txt)
        if len(ln.strip()) >= 4 and sum(c.isalpha() for c in ln) >= 3
    ]
    return "\n".join(linhas)


def _from_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Planilha: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(
        " | ".join(c for c in row if c.strip()) for row in reader if any(row)
    )
